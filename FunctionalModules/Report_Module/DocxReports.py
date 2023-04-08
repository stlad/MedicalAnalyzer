import os

from docx import Document
from copy import deepcopy

from matplotlib import *

from FunctionalModules.DB_Module.db_module import *
from utilits import  date_sql_to_text_format
from Models.Patient import Patient
from Models.Analysis import Analysis
from FunctionalModules.Diagram_Module.Diagram_Processing import DiagramProcessor
from FunctionalModules.Diagnose_Module.Recommendation_Processing import RecommendationProcessor
from docx.shared import Mm

class DocxReporter:

    def __init__(self, analysis:Analysis):
        self.patient = analysis.patient
        self.analysis = analysis

        self.doc = self._get_report()


    def _get_report(self):
        doc = self._get_classic_table()
        doc = self._add_graphs(doc)
        doc = self._add_recommendations(doc)


        return doc

    def _get_table_from_template(self):
        doc = Document('FunctionalModules/Report_Module/form_template.docx')
        #return deepcopy(doc.tables[0])
        return doc.tables[0]


    def _get_classic_table(self):
        doc = Document('FunctionalModules/Report_Module/form_template.docx')

        tbl = doc.tables[0]

        tbl.rows[0].cells[1].text = f"{self.patient.surname} {self.patient.name} {self.patient.patronymic} "
        tbl.rows[1].cells[1].text = self.patient.gender # ПОЛ
        tbl.rows[2].cells[1].text = date_sql_to_text_format(str(self.analysis.analysis_date)) # ДАТА АНАЛИЗА
        tbl.rows[3].cells[1].text = f'{self.patient.get_age_by_date(self.analysis.analysis_date)} Полных лет'# ВОЗРАСТ
        tbl.rows[4].cells[1].text = self.patient.diag # ДИАГНОЗ 1
        tbl.rows[5].cells[1].text = self.patient.second_diag # ДИАГНОЗ 2

        tbl.rows[6].cells[1].text = self.patient.genes  # ГЕНЫ 1,2,3
        tbl.rows[7].cells[1].text = self.patient.genes  # ГЕНЫ 1,2,3
        tbl.rows[8].cells[1].text = self.patient.genes  # ГЕНЫ 1,2,3
        tbl.rows[9].cells[1].text = '0' if self.analysis.analysis_date.month in [3,4,5,6,7,8] else '1'

        self._fill_params(tbl)
        return doc


    def _fill_params(self, tbl):
        i=10
        for param in self.analysis.parameters:
            if i==10 or i == 19 or i == 30 :
                i+=1
            tbl.rows[i].cells[0].text = param.name
            tbl.rows[i].cells[1].text = str(param.value)
            i += 1

    def _add_graphs(self, document :Document):
        d_processor = DiagramProcessor()
        t,b = self._get_radar_graphs(d_processor)
        triangle = self._get_triangle_graph(d_processor)
        linear = self._get_lineargraph(d_processor)
        self._add_figure_to_doc(document, t, 'Т_клеточное звено.png')
        self._add_figure_to_doc(document, b, 'В_клеточное звено.png')
        #self._add_figure_to_doc(document, linear[0], 'ф.png')
        #self._add_figure_to_doc(document, linear[1], 'as.png')
        self._add_figure_to_doc(document, triangle, 'цитокиновые пары.png')

        return document

    def _add_recommendations(self, doc: Document):
        rec_proc = RecommendationProcessor()
        rec = rec_proc.MakeRecommendation(self.patient, self.analysis.analysis_date)
        p = doc.add_paragraph()
        run = p.add_run("Рекомендации: \n\n"+ rec)
        return doc

    def _add_figure_to_doc(self, doc:Document, fig, img_name):
        fig.savefig(img_name)
        p = doc.add_paragraph()
        run = p.add_run('')
        run.add_picture(img_name, height=Mm(150))
        os.remove(img_name)


    def _get_radar_graphs(self, diagram_processor: DiagramProcessor):
        t,b = diagram_processor.MakeRadar(self.patient, self.analysis.analysis_date)
        return t.Figure, b.Figure

    def _get_lineargraph(self, diagram_processor: DiagramProcessor):
        figs = diagram_processor.MakeTimeDiagram(self.patient,
                                                 self.patient.analysis[0].analysis_date,
                                                 self.patient.analysis[-1].analysis_date)
        return figs[0].Figure, figs[1].Figure

    def _get_triangle_graph(self,diagram_processor: DiagramProcessor):
        g = diagram_processor.MakeTriangleDiagram(self.patient, self.analysis.analysis_date)
        return g.Figure




    def save_to_file(self, filename):
        if self.doc == None:
            return
        self.doc.save(filename)
