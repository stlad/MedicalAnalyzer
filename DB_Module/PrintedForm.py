import pandas as pd
from docx import Document
from copy import deepcopy
#from DB_Module.db_module import *
from utilits import  date_sql_to_text_format

class PrintedForm():

    def __init__(self, patient:list, analysis):
        self.patient = patient
        self.analysis = analysis
        t = self._get_table_from_template()
        self._get_classic_table(t)
        pass

    def _get_table_from_template(self):
        doc = Document('form_template')
        return deepcopy(doc.tables[0])


    def _get_classic_table(self, template_table):
        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph._p.addnext(template_table._tbl)

        tbl = doc.tables[0]

        tbl.rows[0].cells[1].text = self.patient[2]+self.patient[1]+self.patient[3] # ФИО
        tbl.rows[1].cells[1].text = self.patient[8] # ПОЛ
        tbl.rows[2].cells[1].text = date_sql_to_text_format(str(self.analysis[2])) # ДАТА АНАЛИЗА
        tbl.rows[3].cells[1].text = str(int((self.analysis[2] - self.patient[4]).days/ 365)) + ' Полных лет'# ВОЗРАСТ
        tbl.rows[4].cells[1].text = self.patient[5] # ДИАГНОЗ 1
        tbl.rows[5].cells[1].text = self.patient[6] # ДИАГНОЗ 2

        tbl.rows[6].cells[1].text = self.patient[7] # ГЕНЫ 1,2,3
        tbl.rows[7].cells[1].text = self.patient[7]  # ГЕНЫ 1,2,3
        tbl.rows[8].cells[1].text = self.patient[7]  # ГЕНЫ 1,2,3

        tbl.rows[9].cells[1].text = '0' if self.analysis[2].month in [3,4,5,6,7,8] else '1'

        doc.save('hello.docx')



    def save_to_file(self, filename):
        pass



'''pat = MainDBController.GetAllPatients()[0]
anal = MainDBController.GetAllAnalysisByPatientID(5)[0]
print(pat)
print(str(anal[2]))
print(pat[4])
PrintedForm(pat,anal)'''


# пациент [ID, имя, фамилия, отчество, дата рождения, диаг, диаг2, гены, пол, возраст]
# анализ [ID, ID_пациента, дата]
# каталог [ID, название, ед, от, до]
# параметр [ID, ID_по_каталогу, значение, ID_анализа, отклонение]
